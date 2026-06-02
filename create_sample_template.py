"""
create_sample_template.py
--------------------------
One-time script to create:
  1. templates/sample_notice.docx  — Word notice template with Jinja2 placeholders
  2. sample/sample_data.xlsx       — Sample Excel file with 5 test rows

Run once before testing:
    py -3 create_sample_template.py
"""

import os

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import openpyxl.styles


BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def create_docx_template() -> str:
    """
    Create a sample legal notice Word template with docxtpl Jinja2 placeholders.
    Uses python-docx to build the document programmatically.
    """
    doc = Document()

    # ── Page margins (narrow) ─────────────────────────────────────────────
    section = doc.sections[0]
    section.left_margin = section.right_margin = 914400      # 1 inch in EMUs
    section.top_margin = section.bottom_margin = 914400

    # ── Header: Firm name ─────────────────────────────────────────────────
    firm_heading = doc.add_paragraph()
    firm_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = firm_heading.add_run("{{ FIRM_NAME }}")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()  # spacer

    # ── Notice title ──────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("LEGAL NOTICE")
    title_run.bold = True
    title_run.underline = True
    title_run.font.size = Pt(13)

    doc.add_paragraph()

    # ── Date and reference ────────────────────────────────────────────────
    doc.add_paragraph("Date: {{ NOTICE_DATE }}")
    doc.add_paragraph("Ref No: {{ ACCOUNTNO }}")

    doc.add_paragraph()

    # ── Addressee ─────────────────────────────────────────────────────────
    doc.add_paragraph("To,")
    doc.add_paragraph("{{ NAME }}")
    doc.add_paragraph()

    # ── Salutation ────────────────────────────────────────────────────────
    doc.add_paragraph("Dear {{ NAME }},")
    doc.add_paragraph()

    # ── Body ──────────────────────────────────────────────────────────────
    body1 = doc.add_paragraph()
    body1.add_run(
        "This notice is being sent to you on behalf of our client in the matter of "
        "your account number "
    )
    body1.add_run("{{ ACCOUNTNO }}").bold = True
    body1.add_run(
        ", wherein a cheque of Rs. "
    )
    body1.add_run("{{ AMOUNT }}").bold = True
    body1.add_run(
        " dated "
    )
    body1.add_run("{{ CHEQUE_DATE }}").bold = True
    body1.add_run(
        " drawn on "
    )
    body1.add_run("{{ BRANCH }}").bold = True
    body1.add_run(
        " has been returned / dishonoured due to: "
    )
    body1.add_run("{{ REASON }}.").bold = True

    doc.add_paragraph()

    doc.add_paragraph(
        "You are hereby called upon to make the payment of the aforesaid amount of "
        "Rs. {{ AMOUNT }} within 30 (thirty) days from the receipt of this notice, "
        "failing which legal proceedings shall be initiated against you without any "
        "further notice."
    )

    doc.add_paragraph()

    doc.add_paragraph(
        "For any queries or to make payment, please contact: {{ OFFICER_NO }}"
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Signature ─────────────────────────────────────────────────────────
    doc.add_paragraph("Yours faithfully,")
    doc.add_paragraph()
    doc.add_paragraph("{{ LAWYER_NAME }}")
    doc.add_paragraph("{{ FIRM_NAME }}")
    doc.add_paragraph("Date: {{ NOTICE_DATE }}")

    # ── Save ──────────────────────────────────────────────────────────────
    templates_dir = os.path.join(BASE_DIR, "templates")
    os.makedirs(templates_dir, exist_ok=True)
    output_path = os.path.join(templates_dir, "sample_notice.docx")
    doc.save(output_path)
    print(f"✅ Template created: {output_path}")
    return output_path


def create_sample_excel() -> str:
    """
    Create a sample Excel file with 5 test rows matching the template variables.
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Notices"

    # ── Headers ───────────────────────────────────────────────────────────
    headers = [
        "NAME", "EMAILID", "MOBILENO", "AMOUNT",
        "ACCOUNTNO", "CHEQUE_DATE", "REASON", "BRANCH", "OFFICER_NO",
    ]
    ws.append(headers)

    # Style header row
    header_font = openpyxl.styles.Font(bold=True)
    header_fill = openpyxl.styles.PatternFill("solid", fgColor="4472C4")
    header_font_white = openpyxl.styles.Font(bold=True, color="FFFFFF")
    for col_idx, _ in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font_white

    # ── Sample data (5 rows) ──────────────────────────────────────────────
    sample_rows = [
        ["Ramesh Kumar",   "ramesh.kumar@example.com",  "9876543210", "50000",
         "HDFC-001-2026", "15/01/2026", "Insufficient Funds",     "HDFC Pune Main",  "8830575674"],
        ["Sunita Patil",   "sunita.patil@example.com",  "8765432109", "25000",
         "HDFC-002-2026", "20/01/2026", "Account Closed",          "HDFC Kothrud",    "8830575674"],
        ["Ajay Mehta",     "ajay.mehta@example.com",    "7654321098", "75000",
         "HDFC-003-2026", "22/01/2026", "Payment Stopped",         "HDFC Baner",      "8830575674"],
        ["Priya Sharma",   "priya.sharma@example.com",  "9988776655", "15000",
         "HDFC-004-2026", "28/01/2026", "Signature Mismatch",      "HDFC Wakad",      "8830575674"],
        ["Vikram Singh",   "vikram.singh@example.com",  "9123456780", "100000",
         "HDFC-005-2026", "30/01/2026", "Insufficient Funds",     "HDFC Shivajinagar","8830575674"],
    ]
    for row in sample_rows:
        ws.append(row)

    # Auto-fit column widths
    for column in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in column)
        ws.column_dimensions[column[0].column_letter].width = max_length + 4

    # ── Save ──────────────────────────────────────────────────────────────
    sample_dir = os.path.join(BASE_DIR, "sample")
    os.makedirs(sample_dir, exist_ok=True)
    output_path = os.path.join(sample_dir, "sample_data.xlsx")
    wb.save(output_path)
    print(f"✅ Sample Excel created: {output_path}")
    return output_path


if __name__ == "__main__":
    print("Creating sample files for BulkNoticeAutomation...\n")

    # Check python-docx is available (separate from docxtpl)
    try:
        from docx import Document as _check
    except ImportError:
        print("Installing python-docx (needed for template creation only)...")
        import subprocess, sys
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        from docx import Document

    template_path = create_docx_template()
    excel_path = create_sample_excel()

    print(f"\n✅ Done! Files created:")
    print(f"   Template : {template_path}")
    print(f"   Excel    : {excel_path}")
    print(f"\nNow run the test pipeline:")
    print(f"   py -3 test_pipeline.py")
